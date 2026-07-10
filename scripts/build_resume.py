#!/usr/bin/env python3
"""Build Aiden Rhaa's ATS-first two-page resume as a Word document.

The final public PDF is produced from this structured DOCX with LibreOffice,
then metadata-normalized and verified by the release workflow.
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


INK = RGBColor(0x11, 0x10, 0x0D)
MUTED = RGBColor(0x4F, 0x4D, 0x47)
SIGNAL_RED = RGBColor(0xC5, 0x2B, 0x22)
RULE = "BBB5A7"
FONT = "Arial"


def set_run_font(
    run,
    *,
    size: float,
    color: RGBColor = INK,
    bold: bool = False,
    italic: bool = False,
    underline: bool | None = None,
) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    if underline is not None:
        run.underline = underline


def set_paragraph_border_bottom(paragraph, color: str = RULE, size: int = 5) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def add_hyperlink(paragraph, text: str, url: str, *, size: float = 8.9) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), FONT)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "11100D")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    font_size = OxmlElement("w:sz")
    font_size.set(qn("w:val"), str(round(size * 2)))
    font_size_cs = OxmlElement("w:szCs")
    font_size_cs.set(qn("w:val"), str(round(size * 2)))
    run_properties.extend([fonts, color, underline, font_size, font_size_cs])
    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_plain_paragraph(
    doc: Document,
    text: str,
    *,
    size: float = 9.5,
    color: RGBColor = INK,
    bold: bool = False,
    italic: bool = False,
    before: float = 0,
    after: float = 2,
    line_spacing: float = 1.04,
    keep_next: bool = False,
) -> object:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing
    if keep_next:
        keep_with_next(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return paragraph


def add_section_heading(doc: Document, text: str) -> object:
    paragraph = doc.add_paragraph(style="ATS Section")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text.upper())
    set_run_font(run, size=10.5, color=SIGNAL_RED, bold=True)
    set_paragraph_border_bottom(paragraph)
    return paragraph


def add_entry_heading(
    doc: Document,
    title: str,
    descriptor: str | None = None,
    *,
    before: float = 2.8,
) -> object:
    paragraph = doc.add_paragraph(style="ATS Entry")
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.keep_with_next = True
    title_run = paragraph.add_run(title)
    set_run_font(title_run, size=9.7, bold=True)
    if descriptor:
        detail_run = paragraph.add_run(f" | {descriptor}")
        set_run_font(detail_run, size=8.9, color=MUTED)
    return paragraph


def add_link_line(doc: Document, label: str, url: str) -> object:
    paragraph = doc.add_paragraph(style="ATS Link")
    paragraph.paragraph_format.keep_with_next = True
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, size=8.9, color=MUTED, bold=True)
    add_hyperlink(paragraph, url, url, size=8.9)
    return paragraph


def add_bullet(doc: Document, text: str, *, after: float = 1.6) -> object:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.30)
    paragraph.paragraph_format.first_line_indent = Inches(-0.16)
    paragraph.paragraph_format.right_indent = Inches(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.03
    paragraph.paragraph_format.widow_control = True
    run = paragraph.add_run(text)
    set_run_font(run, size=9.5)
    return paragraph


def add_skill_line(doc: Document, label: str, text: str) -> None:
    paragraph = doc.add_paragraph(style="ATS Body")
    paragraph.paragraph_format.space_after = Pt(1.2)
    paragraph.paragraph_format.line_spacing = 1.02
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, size=9.25, bold=True)
    value_run = paragraph.add_run(text)
    set_run_font(value_run, size=9.25)


def add_contact_block(doc: Document) -> None:
    name = doc.add_paragraph(style="ATS Name")
    name_run = name.add_run("AIDEN RHAA")
    set_run_font(name_run, size=20, bold=True)

    headline = doc.add_paragraph(style="ATS Headline")
    headline_run = headline.add_run(
        "Cloud Platform & Reliability Engineer | AWS Infrastructure | Terraform | DevSecOps"
    )
    set_run_font(headline_run, size=10.5, bold=True)

    contact = doc.add_paragraph(style="ATS Contact")
    set_run_font(contact.add_run("Atlanta Metropolitan Area | 617-939-1648 | "), size=8.9)
    add_hyperlink(contact, "aidenrhaacloud@gmail.com", "mailto:aidenrhaacloud@gmail.com", size=8.9)

    links = doc.add_paragraph(style="ATS Contact")
    set_run_font(links.add_run("Portfolio: "), size=8.9, color=MUTED, bold=True)
    add_hyperlink(
        links,
        "https://manynames3.github.io/cloudresume/",
        "https://manynames3.github.io/cloudresume/",
        size=8.9,
    )
    set_run_font(links.add_run(" | GitHub: "), size=8.9, color=MUTED, bold=True)
    add_hyperlink(links, "https://github.com/manynames3", "https://github.com/manynames3", size=8.9)

    linkedin = doc.add_paragraph(style="ATS Contact")
    set_run_font(linkedin.add_run("LinkedIn: "), size=8.9, color=MUTED, bold=True)
    add_hyperlink(
        linkedin,
        "https://linkedin.com/in/aidenrhaa",
        "https://linkedin.com/in/aidenrhaa",
        size=8.9,
    )


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.58)
    section.right_margin = Inches(0.58)
    section.top_margin = Inches(0.52)
    section.bottom_margin = Inches(0.52)
    section.header_distance = Inches(0.20)
    section.footer_distance = Inches(0.20)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.04

    styles = doc.styles
    name_style = styles.add_style("ATS Name", WD_STYLE_TYPE.PARAGRAPH)
    name_style.paragraph_format.space_before = Pt(0)
    name_style.paragraph_format.space_after = Pt(0.8)
    name_style.paragraph_format.line_spacing = 1.0
    name_style.paragraph_format.keep_with_next = True

    headline_style = styles.add_style("ATS Headline", WD_STYLE_TYPE.PARAGRAPH)
    headline_style.paragraph_format.space_before = Pt(0)
    headline_style.paragraph_format.space_after = Pt(1.5)
    headline_style.paragraph_format.line_spacing = 1.0
    headline_style.paragraph_format.keep_with_next = True

    contact_style = styles.add_style("ATS Contact", WD_STYLE_TYPE.PARAGRAPH)
    contact_style.paragraph_format.space_before = Pt(0)
    contact_style.paragraph_format.space_after = Pt(0.5)
    contact_style.paragraph_format.line_spacing = 1.0
    contact_style.paragraph_format.keep_with_next = True

    section_style = styles.add_style("ATS Section", WD_STYLE_TYPE.PARAGRAPH)
    section_style.paragraph_format.space_before = Pt(5.0)
    section_style.paragraph_format.space_after = Pt(2.4)
    section_style.paragraph_format.line_spacing = 1.0

    entry_style = styles.add_style("ATS Entry", WD_STYLE_TYPE.PARAGRAPH)
    entry_style.paragraph_format.space_before = Pt(2.8)
    entry_style.paragraph_format.space_after = Pt(0.5)
    entry_style.paragraph_format.line_spacing = 1.0

    link_style = styles.add_style("ATS Link", WD_STYLE_TYPE.PARAGRAPH)
    link_style.paragraph_format.space_before = Pt(0)
    link_style.paragraph_format.space_after = Pt(1.0)
    link_style.paragraph_format.line_spacing = 1.0

    body_style = styles.add_style("ATS Body", WD_STYLE_TYPE.PARAGRAPH)
    body_style.paragraph_format.space_before = Pt(0)
    body_style.paragraph_format.space_after = Pt(2)
    body_style.paragraph_format.line_spacing = 1.04

    list_style = styles["List Bullet"]
    list_style.font.name = FONT
    list_style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    list_style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    list_style.font.size = Pt(9.5)
    list_style.font.color.rgb = INK

    # Named design resolution: compact_reference_guide + ATS_resume override.
    doc.core_properties.title = "Aiden Rhaa - Cloud Platform & Reliability Engineer Resume"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.subject = "Cloud platform and reliability engineering resume"
    doc.core_properties.comments = ""


def build_resume(output_path: Path) -> None:
    doc = Document()
    configure_document(doc)
    add_contact_block(doc)

    add_section_heading(doc, "Professional Summary")
    add_plain_paragraph(
        doc,
        "Cloud platform and reliability engineer who turns operational and customer requirements into secure, reviewable AWS systems. Builds Terraform-managed serverless, container, and Kubernetes platforms with CI/CD, least-privilege IAM, observability, policy gates, cost controls, tests, runbooks, and documented teardown. Brings 15+ years of workflow ownership and client delivery, including 200+ completed client projects, to create reliable handoffs and systems other operators can trust.",
        after=1.8,
    )

    add_section_heading(doc, "Previously Earned Credentials")
    add_plain_paragraph(
        doc,
        "Previously earned credentials (not currently active): AWS Certified Solutions Architect - Associate; AWS Certified Developer - Associate; HashiCorp Certified: Terraform Associate.",
        after=1.6,
    )

    add_section_heading(doc, "Technical Skills")
    add_skill_line(
        doc,
        "AWS",
        "VPC, IAM, ECS Fargate, EKS, Lambda, API Gateway, Step Functions, S3, SQS, DynamoDB, RDS PostgreSQL, RDS Proxy, Bedrock, Cognito, CloudFront, WAF, ALB, ECR, EventBridge, Transcribe, CloudWatch, CloudTrail, Secrets Manager, SSM",
    )
    add_skill_line(
        doc,
        "Infrastructure & delivery",
        "Terraform, CloudFormation exposure, GitHub Actions, OIDC, Docker, Helm, ArgoCD, Checkov, TFLint, policy gates, CI/CD, runbooks, ADRs, documented teardown",
    )
    add_skill_line(
        doc,
        "Containers & operations",
        "Kubernetes, ECS Fargate, Amazon EKS, IRSA, External Secrets, Prometheus, Grafana, OpenTelemetry, alarms, audit logging, incident and recovery workflows",
    )
    add_skill_line(
        doc,
        "Programming & customer systems",
        "Python, TypeScript, FastAPI, PostgreSQL, SQL, REST APIs, CRM/API automation, webhooks, voice AI guardrails, asynchronous transcription, human review workflows",
    )

    add_section_heading(doc, "Selected Cloud Platform Engineering")
    add_entry_heading(doc, "Clearpath - AWS Fargate Workflow API", "Independent reference system")
    add_link_line(doc, "Repository", "https://github.com/manynames3/clearpath-fargate-api")
    add_bullet(
        doc,
        "During a short-lived May 2026 AWS validation run, deployed a Terraform-managed ECS Fargate stack with two healthy tasks, healthy API and webhook target groups, private RDS PostgreSQL through RDS Proxy, CloudFront /health behind WAF, and visible CloudWatch alarms.",
    )
    add_bullet(
        doc,
        "Latest local validation recorded 36 passing FastAPI tests and 339 passing Terraform Checkov checks (0 failed, 44 skipped); after evidence capture, Terraform destroy removed 93 managed resources to avoid idle cost.",
    )
    add_bullet(
        doc,
        "Translated lead-intelligence operations into REST APIs, GHL-compatible webhook intake, source ROI and funnel analytics, county resolution, dashboard workflows, and documented runbooks.",
    )

    add_entry_heading(doc, "TerraGate - Terraform PR Risk Gate", "Live constrained demo")
    add_link_line(doc, "Repository", "https://github.com/manynames3/terragate")
    add_bullet(
        doc,
        "Built deterministic Python security, cost, reliability, and governance checks with sensitive-value redaction, audit records, evidence-linked remediation, and human approval gates.",
    )
    add_bullet(
        doc,
        "Public demo keeps external writes mocked and Terraform execution disabled, making the boundary inspectable without implying infrastructure mutation.",
    )

    add_entry_heading(doc, "InspectIQ - Automotive Inspection Review", "Live AWS-backed, read-only walkthrough")
    add_link_line(doc, "Repository", "https://github.com/manynames3/inspectiq")
    add_bullet(
        doc,
        "Built a vehicle-photo inspection workflow with private S3, queued Bedrock analysis, app-level Cognito JWT/RBAC, deterministic grading, audit records, and human reviewer accept, reject, and edit gates.",
    )
    add_bullet(
        doc,
        "The public walkthrough is read-only; evidence does not claim API Gateway authorizer enforcement or production model accuracy from the local deterministic evaluation harness.",
    )

    add_entry_heading(doc, "Pulpit V2 - EKS / GitOps Platform", "Validated and torn down")
    add_link_line(doc, "Repository", "https://github.com/manynames3/pulpit-v2")
    add_bullet(
        doc,
        "Validated reusable Terraform, ECR, Helm, ArgoCD, tenant namespaces, External Secrets/SSM, Prometheus, Grafana, ALB ingress, runtime captures, and teardown artifacts.",
        after=0,
    )

    doc.add_page_break()

    add_section_heading(doc, "Professional Experience")
    add_entry_heading(
        doc,
        "Founder / Automation & Systems Lead - Clearpath Property Group / Boston Probate Solutions",
        "Dec 2016 - Present | Atlanta Metropolitan Area",
        before=0,
    )
    add_bullet(
        doc,
        "Own customer-engagement and operations workflows spanning lead intake, CRM integration, seller communication, public-record research, market data, and technical automation.",
    )
    add_bullet(
        doc,
        "Developed Python ingestion pipelines and webhook-driven CRM/API workflows for normalized county records, structured handoff, intent capture, transcript summaries, and voice AI guardrails.",
    )
    add_bullet(
        doc,
        "Translated operating requirements into the public Clearpath reference system with architecture, tests, AWS validation artifacts, runbooks, limitations, and teardown evidence.",
    )

    add_entry_heading(
        doc,
        "Creative Director / Operations - Aiden Rhaa Photography / Visual Impact Studios",
        "Oct 2010 - Present | Atlanta Metropolitan Area",
    )
    add_bullet(
        doc,
        "Delivered 200+ client projects across 15+ years while coordinating contractors, vendors, budgets, timelines, technical execution, and high-touch client communication.",
    )
    add_bullet(
        doc,
        "Managed client-facing web properties and deployment workflows across DNS, SSL/TLS, Cloudflare, hosting platforms, GitHub deployments, cPanel, SFTP/SSH, WordPress, Supabase, and Netlify.",
    )

    add_entry_heading(
        doc,
        "Multimedia Producer / Technical Workflow Lead - ProMedia Productions",
        "May 2009 - Dec 2016 | Boston, MA",
    )
    add_bullet(
        doc,
        "Managed technical production workflows for corporate, education, and small-business media projects, coordinating recording, post-production, delivery, and client handoff.",
    )
    add_bullet(
        doc,
        "Automated repetitive media-file processes and designed custom recording-studio configurations through structured workflow design and tooling.",
    )

    add_section_heading(doc, "Additional Public Systems")
    add_entry_heading(doc, "Super Transcriber", "Serverless audio and voice workflow", before=0)
    add_link_line(doc, "Repository", "https://github.com/manynames3/super-transcriber-api")
    add_bullet(
        doc,
        "Built a cost-first event-driven transcription path with authenticated intake, queued workers, Amazon Transcribe events, webhook delivery, OpenAPI, an SDK, infrastructure as code, and explicit operating limits.",
    )

    add_entry_heading(doc, "AegisDesk", "Cloud operations control plane")
    add_link_line(doc, "Repository", "https://github.com/manynames3/aegisdesk-cloudops-control-plane")
    add_bullet(
        doc,
        "Standardized incident triage, access, governance, and cost-review workflows with role-aware decisions, request replay, answer provenance, redaction, audit trails, throttling, and documented AWS deployment paths.",
    )

    add_entry_heading(doc, "Pulpit", "Governed serverless conversational search")
    add_link_line(doc, "Repository", "https://github.com/manynames3/pulpit")
    add_bullet(
        doc,
        "Built cited Korean-English retrieval with Bedrock, Lambda, API Gateway, Cognito, DynamoDB, private storage, guardrails, audit records, Terraform, and GitHub Actions.",
    )

    add_entry_heading(doc, "PhotoScribe AI", "Governed serverless media search")
    add_link_line(doc, "Repository", "https://github.com/manynames3/photoscribe-ai")
    add_bullet(
        doc,
        "Built private media intake and search with policy and audit records, Bedrock-assisted indexing, queued failure handling, Cognito, DynamoDB, CloudWatch alarms, Terraform, and documented cost boundaries.",
    )

    add_entry_heading(doc, "DocuFlow OCR", "Serverless document-processing workflow")
    add_link_line(doc, "Repository", "https://github.com/manynames3/docuflow-ocr")
    add_bullet(
        doc,
        "Orchestrated presigned S3 intake, Step Functions, Textract, Python Lambda parsing and scoring, DynamoDB job and audit records, SQS dead-letter handling, human review, alarms, and Terraform.",
    )

    add_entry_heading(doc, "FaceID", "Private event gallery with serverless face matching")
    add_link_line(doc, "Repository", "https://github.com/manynames3/faceid")
    add_bullet(
        doc,
        "Built a Cognito-protected gallery with private S3, Rekognition, DynamoDB, Terraform, consent-attested guest intake, owner-scoped delete flows, and human review gates.",
    )

    add_section_heading(doc, "Education")
    add_entry_heading(doc, "Berklee College of Music", "Music / Songwriting | 2008 - 2011", before=0)
    add_entry_heading(doc, "Saint Louis University", "Engineering Physics | 2006 - 2008", before=1.4)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--out",
        type=Path,
        default=Path("work/resume/Aiden_Rhaa_Cloud_Platform_Engineer_Resume.docx"),
    )
    args = parser.parse_args()
    build_resume(args.out)
    print(args.out)


if __name__ == "__main__":
    main()
